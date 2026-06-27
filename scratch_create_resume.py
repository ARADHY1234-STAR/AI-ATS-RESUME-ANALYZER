import docx

doc = docx.Document()
doc.add_heading('John Doe', 0)
doc.add_paragraph('Email: john.doe@example.com | Phone: 123-456-7890 | Address: 123 Main St, Springfield, IL 62701')

doc.add_heading('Professional Summary', level=1)
doc.add_paragraph('Experienced software engineer with a track record of building web applications.')

doc.add_heading('Experience', level=1)
p1 = doc.add_paragraph()
p1.add_run('Software Engineer at TechCorp (2020 - Present)\n').bold = True
doc.add_paragraph('• Responsible for developing new features using Python and React.')
doc.add_paragraph('• Worked on improving site reliability.')

doc.add_heading('Education', level=1)
doc.add_paragraph('Bachelor of Science in Computer Science')

doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, React, SQL, Java, Docker, Kubernetes, AWS, C++')

doc.save('test_resume.docx')
print("Resume created successfully!")
